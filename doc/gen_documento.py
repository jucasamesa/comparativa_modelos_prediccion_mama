# -*- coding: utf-8 -*-
"""Genera el documento tecnico final (>=30 pag) del proyecto de prediccion prospectiva de
cancer de mama EPS Sanitas. Integra creacion de datasets, comparativa de modelos, discusion
de por que XGBoost gano, analisis SHAP por modelo, figuras del proyecto, razon por la que las
flags no afectan, y diseno operativo + A/B. Regla de integridad: toda cifra es trazable
(memoria del proyecto + artefactos JSON/CSV). Figuras en _doc/figs/."""
import os, json
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

BASE = os.path.dirname(os.path.abspath(__file__))
FIGS = os.path.join(BASE, "_doc", "figs")
OUT_PATH = os.path.join(BASE, "_doc", "Documento_tecnico_prediccion_cancer_mama.docx")

doc = Document()
style_normal = doc.styles["Normal"]
style_normal.font.name = "Times New Roman"
style_normal.font.size = Pt(12)
for h_level in ("Heading 1", "Heading 2", "Heading 3"):
    s = doc.styles[h_level]
    s.font.name = "Times New Roman"
    s.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
for section in doc.sections:
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3); section.right_margin = Cm(2.5)

_h1seen = [False]
def h1(t):
    if _h1seen[0]:
        doc.add_page_break()
    _h1seen[0] = True
    doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)
def blank(): doc.add_paragraph()

def p(text, bold_parts=None):
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_parts is None:
        r = par.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(12)
    else:
        remaining = text
        for part in bold_parts:
            idx = remaining.find(part)
            if idx == -1: continue
            before = remaining[:idx]
            if before:
                r = par.add_run(before); r.font.name = "Times New Roman"; r.font.size = Pt(12)
            r = par.add_run(part); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(12)
            remaining = remaining[idx+len(part):]
        if remaining:
            r = par.add_run(remaining); r.font.name = "Times New Roman"; r.font.size = Pt(12)
    return par

def bullet(text, bold_lead=None):
    par = doc.add_paragraph(style="List Bullet"); par.paragraph_format.left_indent = Cm(1)
    if bold_lead:
        r = par.add_run(bold_lead); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(12)
        r2 = par.add_run(text); r2.font.name = "Times New Roman"; r2.font.size = Pt(12)
    else:
        r = par.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(12)

def add_table(headers, rows, widths=None, font=10):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = "Table Grid"
    hc = t.rows[0].cells
    for i, h in enumerate(headers):
        hc[i].text = str(h)
        for run in hc[i].paragraphs[0].runs:
            run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(font)
        hc[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for rd in rows:
        rc = t.add_row().cells
        for i, v in enumerate(rd):
            rc[i].text = str(v)
            for run in rc[i].paragraphs[0].runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(font)
            rc[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    blank()
    return t

_fign = [0]
def figure(fname, caption, width=14.5):
    path = os.path.join(FIGS, fname)
    if not os.path.exists(path):
        p(f"[FIGURA NO ENCONTRADA: {fname}]"); return
    _fign[0] += 1
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(path, width=Cm(width))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figura {_fign[0]}. {caption}")
    r.italic = True; r.font.name = "Times New Roman"; r.font.size = Pt(10)
    blank()

def title(text, size=20, color=(0x1F,0x38,0x64), after=6):
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(text); r.bold = True; r.font.name = "Times New Roman"
    r.font.size = Pt(size); r.font.color.rgb = RGBColor(*color)
    par.paragraph_format.space_after = Pt(after)

# ============================ PORTADA ============================
for _ in range(3): blank()
title("Predicción prospectiva de incidencia de cáncer de mama en EPS Sanitas", 22)
title("Construcción de datasets, comparativa de modelos de aprendizaje automático e "
      "interpretabilidad con SHAP para un programa de tamización proactiva", 13, (0x40,0x40,0x40), 18)
for _ in range(2): blank()
title("Documento técnico del proyecto", 14, (0,0,0), 4)
title("Cuenta de Alto Costo · Dirección de Gestión del Riesgo", 12, (0x40,0x40,0x40), 4)
for _ in range(6): blank()
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pp.add_run("Enfoque de incidencia (predicción prospectiva) · Modelo ganador: XGBoost · "
               "Validación temporal 2023→2024")
r.italic = True; r.font.size = Pt(11)
doc.add_page_break()

# ============================ RESUMEN ============================
h1("Resumen")
p("Este documento presenta el desarrollo técnico de un modelo de predicción prospectiva de "
  "incidencia de cáncer de mama para la población femenina afiliada a EPS Sanitas. A diferencia "
  "de los modelos de diagnóstico sobre imagen o de los modelos clásicos de evaluación de riesgo "
  "individual (Gail, 1989; Tyrer et al., 2004), el objetivo aquí es anticipatorio: identificar, "
  "entre las mujeres actualmente sin diagnóstico, a quienes presentan mayor probabilidad de "
  "desarrollar un primer cáncer de mama en el año siguiente, para orientar una intervención "
  "proactiva de tamización evaluada mediante un ensayo A/B. El problema se formaliza como una "
  "clasificación binaria con desbalanceo extremo: a partir de pares temporales T→T+1 construidos "
  "sobre 12,4 millones de registros consolidados de cuatro fuentes (radicados de la Cuenta de "
  "Alto Costo, población de afiliados, historia clínica electrónica Avicena y calidad del aire "
  "PM10 del IDEAM), se obtuvieron 1.857 positivos frente a 5,47 millones de negativos en "
  "entrenamiento (tasa de evento 0,034%; razón ≈ 2.945:1).")
p("Se compararon regresión logística, LightGBM y XGBoost bajo distintas estrategias de "
  "tratamiento del desbalanceo (class_weight, SMOTE y reponderación por scale_pos_weight), con "
  "validación cruzada GroupKFold por paciente y una validación temporal final 2023→2024 nunca "
  "vista durante el ajuste. El modelo ganador fue un XGBoost tuneado con Optuna sobre la variante "
  "de datos con NaN nativo (AUC-PR 0,0380; recall@top10% 0,704; AUC-ROC 0,900). Un hallazgo "
  "metodológico central es que la reponderación agresiva degrada el ordenamiento de riesgo: el "
  "óptimo se alcanza con reponderación nula o leve. El análisis SHAP confirmó que el modelo se "
  "apoya principalmente en la edad y la intensidad de uso del sistema de salud, lo que motivó "
  "una depuración de características y un diseño de asignación estratificado por regional para "
  "atender la equidad geográfica. Finalmente se define el punto de operación del programa y se "
  "dimensiona el ensayo A/B sobre la subpoblación menor de 45 años, donde el modelo aporta valor "
  "incremental respecto al tamizaje bienal por edad.")

# ============================ INDICE ============================
doc.add_page_break()
h2("Contenido")
_toc = [
 "Resumen", "1. Introducción", "2. Estado del arte",
 "3. Materiales y métodos", "4. Resultados y discusión",
 "5. Diseño operativo y ensayo A/B", "6. Conclusiones", "7. Referencias",
 "8. Código fuente y datos analizados",
]
for it in _toc:
    par = doc.add_paragraph(); par.paragraph_format.left_indent = Cm(0.5)
    r = par.add_run(it); r.font.name = "Times New Roman"; r.font.size = Pt(12)

# ============================ 1. INTRODUCCION ============================
h1("1. Introducción")
p("El cáncer de mama es la neoplasia más frecuente en mujeres a escala mundial. Según las "
  "estimaciones de GLOBOCAN, en 2020 se registraron 2.261.419 casos nuevos, equivalentes al "
  "11,7% de todos los cánceres diagnosticados (Sung et al., 2021). En Colombia, la Cuenta de "
  "Alto Costo (CAC, 2024) reporta que el cáncer de mama es la primera causa de mortalidad "
  "oncológica en mujeres, con un 22,48% de las muertes por cáncer, y contabiliza 8.590 casos "
  "nuevos invasivos en 2024. La detección tardía sigue siendo un problema estructural: el 28,16% "
  "de los casos nuevos se diagnostican en estadios III–IV, y la mediana del tiempo entre la "
  "sospecha y el inicio del tratamiento es de 92,40 días (DE ±71,02). Estas cifras delimitan una "
  "ventana de oportunidad clara para las estrategias de detección temprana.")
p("La importancia clínica y económica de adelantar el diagnóstico está bien documentada. La "
  "supervivencia a cinco años en estadio I alcanza ≥99% en los subtipos HR+/ERBB2- y ≥85% en el "
  "triple negativo, mientras que el estadio IV triple negativo presenta una supervivencia mediana "
  "de apenas 10–13 meses (Waks & Winer, 2019). En términos de costo, Mariotto et al. (2011) "
  "documentan que la fase inicial de atención en menores de 65 años cuesta USD 27.693 frente a "
  "USD 94.284 en el último año de vida, una razón de 3,4 veces. En conjunto, estos resultados "
  "justifican económicamente y clínicamente cualquier intervención que desplace los diagnósticos "
  "hacia estadios más tempranos. EPS Sanitas, con una población femenina afiliada superior a tres "
  "millones y un registro anual de entre 978 y 1.326 casos nuevos de cáncer de mama entre 2021 y "
  "2024, dispone tanto de la escala poblacional como de los datos administrativos y clínicos "
  "necesarios para construir un sistema de identificación proactiva de riesgo.")
p("El presente trabajo aborda el problema desde un enfoque de incidencia y no de diagnóstico. "
  "Mientras que buena parte de la literatura de aprendizaje automático en cáncer de mama se centra "
  "en clasificar lesiones ya existentes —por ejemplo, los modelos de aprendizaje profundo sobre "
  "mamografía alcanzan sensibilidades del 88% y AUC de 0,92 (Tahir & Khan, 2025), o los ensambles "
  "sobre datos tabulares clínicos donde Random Forest supera a Naive Bayes, KNN y árboles "
  "individuales (Yavuz et al., 2023)— el objetivo de este proyecto es anticipar la aparición de "
  "la enfermedad en población sana. Esta diferencia de planteamiento condiciona todo el diseño: "
  "la variable objetivo es el primer diagnóstico en el año siguiente, la población de referencia "
  "excluye a las pacientes prevalentes y la métrica principal es el recall, porque el costo de un "
  "falso negativo (una mujer de alto riesgo no tamizada) es clínicamente asimétrico respecto al de "
  "un falso positivo.")
p("Conviene además situar la oportunidad de mejora en su contexto histórico. La detección en "
  "estadios tempranos I–IIa en Colombia pasó del 51% en 2019 al 54% en 2024 (CAC, 2024), una "
  "mejora real pero modesta que deja un margen amplio para estrategias de identificación proactiva "
  "del riesgo. En la regional de Bogotá, la prevalencia de período de casos nuevos reportados "
  "(PCNR) alcanza 34,41 por cada 100.000 mujeres. Un sistema capaz de concentrar la tamización en "
  "las mujeres de mayor riesgo permitiría, en principio, acelerar ese desplazamiento de estadios "
  "sin multiplicar el número de estudios realizados, optimizando el uso de la capacidad instalada.")
p("La contribución de este documento es triple. Primero, describe la construcción conceptual y "
  "reproducible de una base de datos de entrenamiento de escala poblacional con control estricto "
  "de fugas de información temporal. Segundo, presenta una comparativa rigurosa de modelos de "
  "aprendizaje automático bajo desbalanceo extremo, con una discusión de las razones por las que "
  "el gradient boosting supera a los modelos lineales y de por qué la reponderación de clases, "
  "lejos de ayudar, degrada el ordenamiento de riesgo. Tercero, aplica SHAP para interpretar y "
  "depurar el modelo, identifica y discute el sesgo de vigilancia, y traduce el modelo en un "
  "diseño operativo de tamización con consideraciones de equidad geográfica y un dimensionamiento "
  "formal del ensayo A/B.")

# ============================ 2. ESTADO DEL ARTE ============================
h1("2. Estado del arte")
h2("2.1. Modelos clásicos de riesgo")
p("Los modelos de Gail (Gail et al., 1989) y Tyrer-Cuzick (Tyrer et al., 2004) son la referencia "
  "histórica en la estimación de riesgo individual de cáncer de mama. Ambos fueron concebidos para "
  "su uso en consulta y se apoyan en variables como la edad de menarquia, la edad al primer parto, "
  "el número de biopsias previas y la historia familiar detallada de primer grado. El modelo "
  "Tyrer-Cuzick incorpora además componentes genéticos y hormonales. Su principal limitación en el "
  "contexto de este proyecto es que dependen de variables que rara vez están completas en registros "
  "administrativos: la historia familiar estructurada, la densidad mamaria y los resultados "
  "detallados de biopsia no se capturan de forma sistemática en las fuentes disponibles. Conviene "
  "precisar que, contrariamente a una creencia extendida, el modelo de Gail (1989) no exige la "
  "densidad mamaria, y el de Tyrer-Cuzick (2004) tampoco la requiere como variable obligatoria.")
h2("2.2. Aprendizaje automático supervisado en cáncer de mama")
p("La literatura reciente muestra un buen desempeño de los métodos de ensamble sobre datos "
  "tabulares. Yavuz et al. (2023) reportan que Random Forest supera consistentemente a Naive Bayes, "
  "KNN y árboles de decisión individuales, tanto en escenarios balanceados como desbalanceados. "
  "Singh & Dubey (2024) revisan el uso de gradient boosting y citan, como referencia secundaria de "
  "Nemande & Fegade (2023), valores de AUC muy elevados para XGBoost en conjuntos de diagnóstico; "
  "es importante señalar que dichas cifras provienen de tareas de diagnóstico sobre datos "
  "balanceados y no son extrapolables a un problema de incidencia con desbalanceo de 3.000:1. En "
  "el plano del aprendizaje profundo aplicado a imagen, Tahir & Khan (2025) reportan sensibilidad "
  "del 88% (IC 95%: 85–90%) y AUC de 0,92. Finalmente, AlGhunaim & Al-Baity (2019) muestran que "
  "las máquinas de soporte vectorial pueden superar a Random Forest y árboles en datos genómicos "
  "de alta dimensionalidad ejecutados sobre Apache Spark, aunque su escalabilidad con millones de "
  "registros y kernels no lineales es problemática.")
h2("2.3. Aprendizaje con clases desbalanceadas")
p("El desbalanceo extremo es el reto técnico dominante de este problema. Las estrategias estándar "
  "incluyen la reponderación de clases (class_weight inversamente proporcional a la frecuencia), "
  "el sobremuestreo sintético de la minoría mediante SMOTE (Chawla et al., 2002) y la optimización "
  "del umbral de decisión a posteriori. Estas técnicas modifican aspectos distintos del problema: "
  "la reponderación y SMOTE alteran la función de pérdida durante el entrenamiento, mientras que la "
  "optimización de umbral solo reubica el punto de corte sin cambiar el ordenamiento de las "
  "puntuaciones. Como se discutirá en los resultados, esta distinción es decisiva cuando la métrica "
  "de interés es el AUC-PR o el recall@top-k, que dependen del ordenamiento y no del umbral.")
p("La distinción entre técnicas que actúan sobre la pérdida y técnicas que actúan sobre el umbral "
  "es especialmente relevante en programas de tamización, donde el número de cupos disponibles fija "
  "de antemano la fracción de población que puede intervenirse. En ese escenario, la decisión "
  "operativa equivale a seleccionar un percentil de riesgo, por lo que lo único que importa del "
  "modelo es que ordene correctamente a las pacientes; el valor absoluto de la probabilidad y la "
  "posición del umbral por defecto se vuelven secundarios. Esta observación, desarrollada en los "
  "resultados, explica por qué métricas como el AUC-PR y el recall@top-k son preferibles a la "
  "exactitud o al F1 calculado en el umbral 0,5, y por qué la reponderación extrema resultó "
  "contraproducente.")
h2("2.4. Interpretabilidad: SHAP")
p("SHAP (SHapley Additive exPlanations; Lundberg & Lee, 2017) es el marco de interpretabilidad "
  "adoptado en este trabajo. Su fundamento son los valores de Shapley de la teoría de juegos "
  "cooperativos, y constituye el único método de atribución aditiva que satisface simultáneamente "
  "tres propiedades deseables: precisión local (los valores SHAP suman exactamente la predicción "
  "del modelo), ausencia (missingness: una variable sin información recibe atribución nula) y "
  "consistencia (si un modelo incrementa su dependencia de una variable, su atribución no "
  "disminuye). En el contexto regulatorio colombiano y para la validación clínica por comités "
  "médicos, la interpretabilidad es un requisito y no un complemento opcional.")

# ============================ 3. MATERIALES Y METODOS ============================
h1("3. Materiales y métodos")
h2("3.1. Fuentes de datos")
p("El sistema integra cuatro fuentes heterogéneas, todas residentes en Google Cloud Storage "
  "(en el entorno de nube institucional) salvo la de calidad del aire, que es local:")
add_table(["Fuente", "Cobertura", "Rol en el modelo", "Variables clave"],
    [["Radicados CAC (post-auditoría)", "2021–2024", "Etiqueta (label) y antecedente oncológico",
      "key, año, cancer_mama, incidentes, CIE-10"],
     ["Población de afiliados", "2021–2024 (registro de diciembre)", "Universo de negativos y demografía",
      "key, sexo, fec_nacim, FechaAfil EPS, cod_municip, ZONA, regional"],
     ["Historia clínica Avicena", "2021–2024", "Variables clínicas, labs y comorbilidades",
      "Antropometría, labs, CIE-10, remisiones, tabaquismo"],
     ["Calidad del aire PM10 (IDEAM)", "2023–2024 (2021–22 extrapolados)", "Feature ambiental por municipio",
      "Código municipio DANE, PM10 µg/m³ anual"]],
    widths=[3.5, 3.2, 4.2, 4.8], font=9)
p("La etiqueta gold standard proviene exclusivamente de los radicados CAC post-auditoría, que "
  "representan diagnósticos validados ante la Cuenta de Alto Costo. La detección de mama en "
  "Avicena sin radicado CAC se usa únicamente para excluir prevalentes, nunca como fuente de "
  "positivos. Para la identificación del cáncer de mama se emplea el conjunto de códigos CIE-10 "
  "C50.0–C50.9 (invasivo) y D05.0–D05.9 (in situ).")

p("La fuente de calidad del aire merece una nota metodológica. El IDEAM solo publica de forma "
  "consistente el contaminante PM10, agregado por estación de monitoreo. El feature municipal se "
  "construye promediando las estaciones por código de municipio y año. Para 2021 y 2022, sin datos "
  "directos, se extrapola linealmente hacia atrás desde la tendencia 2023→2024, con piso en el "
  "mínimo histórico observado; los municipios con un solo año disponible reciben el valor plano. "
  "Se acompaña de un indicador pm10_disponible que vale cero en municipios sin estación, de modo "
  "que el modelo distingue entre 'sin medición' y 'PM10 medido bajo'. La cobertura abarca 35 "
  "municipios con tendencia real y 27 con un único año.")

h2("3.2. Construcción conceptual del conjunto de entrenamiento")
p("El diseño sigue el principio de predicción prospectiva: las características del año T predicen "
  "el primer diagnóstico en el año T+1. Esta separación temporal es la garantía fundamental contra "
  "la fuga de información. Se construyeron tres pares temporales, dos para entrenamiento y uno para "
  "validación, reservando 2024→2025 para producción:")
add_table(["Par temporal", "Split", "Uso"],
    [["2021 → 2022", "Train (par 1)", "Entrenamiento"],
     ["2022 → 2023", "Train (par 2)", "Entrenamiento"],
     ["2023 → 2024", "Validación", "Evaluación final prospectiva"]],
    widths=[4, 4, 7])
p("La distinción entre incidencia y prevalencia es el eje conceptual del diseño y conviene "
  "precisarla. Un modelo de prevalencia respondería a la pregunta '¿quién tiene cáncer ahora?'; el "
  "modelo de incidencia responde a '¿quién, estando sana hoy, lo desarrollará el próximo año?'. "
  "Esta diferencia no es semántica: determina que las pacientes ya diagnosticadas deban excluirse "
  "del universo (no pueden ser un caso incidente futuro) y que ninguna variable del año T+1 pueda "
  "usarse como predictor. La identificación del cáncer de mama emplea el siguiente conjunto de "
  "códigos CIE-10, donde los códigos in situ (D05) tienen máximo valor clínico por representar la "
  "detección más temprana posible:")
add_table(["Categoría", "Códigos CIE-10"],
    [["Carcinoma invasivo de mama", "C50.0, C50.1, C50.2, C50.3, C50.4, C50.5, C50.6, C50.8, C50.9"],
     ["Carcinoma in situ de mama", "D05.0, D05.1, D05.7, D05.9"]],
    widths=[5.5, 9], font=10)
p("La definición de clases sigue una lógica de incidencia estricta. Son positivos (label = 1) las "
  "mujeres presentes en la población del año T, sin diagnóstico de cáncer de mama en T, con primer "
  "diagnóstico confirmado en T+1. Son negativos (label = 0) las mujeres en población T, sin "
  "diagnóstico en T ni en T+1. Se excluyen las prevalentes en T (ya diagnosticadas en CAC o "
  "Avicena) y las mujeres sin continuidad de afiliación entre T y T+1. Tras aplicar estas reglas "
  "sobre la base acumulada femenina de 12.418.059 filas, los resultados de la construcción "
  "(ejecutada el 18/05/2026) son los siguientes:")
add_table(["Par", "Split", "Pob. T (sin mama)", "Continuidad T→T+1", "Positivos", "Negativos", "Event rate"],
    [["2021→2022", "Train", "2.462.241", "2.446.215", "817", "2.445.398", "0,033%"],
     ["2022→2023", "Train", "3.253.092", "3.025.131", "1.040", "3.024.091", "0,034%"],
     ["2023→2024", "Validación", "3.329.988", "3.058.762", "993", "3.057.769", "0,032%"]],
    widths=[2.4, 2.2, 2.8, 2.8, 1.9, 2.5, 1.9], font=9)
p("En conjunto, el entrenamiento reúne 1.857 positivos frente a 5.469.489 negativos (razón ≈ "
  "2.945:1) y la validación 993 positivos frente a 3.057.769 negativos (razón ≈ 3.079:1). El "
  "identificador de paciente (key) se anonimizó con los primeros 16 caracteres del hash SHA-256 "
  "de la concatenación del tipo y número de documento; al ser determinista, preserva el join "
  "temporal y la partición por paciente, con probabilidad de colisión despreciable.")
figure("01_embudo_dataset.png", "Embudo de construcción del conjunto de entrenamiento (escala "
       "logarítmica). La caída desde la población sin diagnóstico de mama hasta los 1.857 positivos "
       "ilustra la magnitud del desbalanceo de clases.", 12)

h2("3.3. Procesamiento de variables clínicas (Avicena)")
p("Para cada paciente-año se selecciona el valor no nulo más reciente de cada variable dentro del "
  "año mediante funciones de ventana en PySpark (last con ignorenulls sobre ventana ascendente por "
  "fecha de apertura de historia clínica, seguido de deduplicación por row_number). Se aplican "
  "filtros de validez fisiológica antes de la agregación: peso y talla se anulan ante valores "
  "imposibles, la talla se normaliza a centímetros y se acota, y las tensiones arteriales se "
  "restringen a rangos plausibles. Los valores inválidos se convierten a nulo sin eliminar el "
  "registro, preservando la señal de que la paciente asistió a consulta.")
p("Para cada laboratorio (hemoglobina glicosilada, colesterol HDL, LDL y total, triglicéridos) se "
  "computa el valor más reciente y el delta intraanual, además de un indicador binario "
  "{lab}_disponible que vale 1 si existe al menos un registro del laboratorio en el año. Las "
  "comorbilidades (hipertensión, diabetes, biopsia benigna previa, cáncer no-mama y cáncer de "
  "mama) se construyen mediante un OR lógico sobre los folios del año, distinguiendo antecedente, "
  "diagnóstico confirmado y consolidado. Se añaden flags de remisión a oncología, ginecología, "
  "mama y endocrinología, y el conteo de consultas anuales (n_consultas_T).")

h2("3.4. Cohorte, análisis exploratorio e imputación")
p("Se fijó una cohorte de edad ≥ 18 años antes de cualquier imputación, lo que elimina 1,18 "
  "millones de menores sin perder ningún positivo y eleva la tasa de evento a 0,0433% en "
  "entrenamiento y 0,0412% en validación. Las medianas de imputación se calcularon únicamente "
  "sobre el train y se aplicaron en validación para evitar fuga. La auditoría de missingness "
  "definió tres regímenes: imputación simple por debajo del 30% de ausencia, imputación más "
  "indicador binario entre 30% y 60%, y exclusión o solo indicador por encima del 60%. Los flags "
  "de Avicena y los deltas de laboratorio se imputan a cero por su semántica (la ausencia indica "
  "que no hubo visita o medición, no un dato faltante aleatorio).")

p("Las medianas de imputación, calculadas sobre la cohorte de entrenamiento ≥ 18 años, se "
  "documentan a continuación para garantizar la reproducibilidad y su aplicación idéntica en "
  "validación y producción. El filtrado por edad antes de imputar desplazó al alza las variables "
  "antropométricas (el IMC pasó de 25,15 a 26,04 y el perímetro abdominal de 84 a 87), mientras "
  "que los laboratorios apenas variaron, pues los menores rara vez los tienen registrados.")
add_table(["Variable", "Mediana (train ≥18)", "Variable", "Mediana (train ≥18)"],
    [["Colesterol total", "189", "IMC", "26,04"],
     ["Colesterol HDL", "48,2", "Talla (cm)", "158"],
     ["Colesterol LDL", "108", "Perímetro abdominal", "87"],
     ["Triglicéridos", "137", "Menarquia", "12"],
     ["Hemoglobina glicosilada", "6,2", "TA sistólica", "116"],
     ["TA diastólica", "70", "ZONA (moda)", "Urbano"]],
    widths=[4, 3.5, 4, 3.5], font=10)
p("El análisis exploratorio identificó además los predictores individualmente más fuertes, medidos "
  "por el tamaño del efecto de Cohen para las variables continuas y por la razón de odds (OR) para "
  "las binarias, todos calculados sobre el conjunto de entrenamiento:")
add_table(["Predictor", "Tipo", "Medida de efecto"],
    [["Edad", "Continua", "Cohen's d = 1,21"],
     ["Número de consultas en T", "Continua", "Cohen's d = 1,14"],
     ["Tiene registro Avicena", "Binaria", "OR ≈ 16"],
     ["Biopsia / antecedente de biopsia", "Binaria", "OR ≈ 10,7"],
     ["Proxy de menopausia (edad ≥ 50)", "Binaria", "OR ≈ 6,2"],
     ["Antecedente familiar de mama", "Binaria", "OR ≈ 5,3"]],
    widths=[6, 3, 4.5], font=10)
p("Es importante anticipar una aparente paradoja que se resolverá con SHAP: aunque el antecedente "
  "familiar de mama presenta una OR elevada en el análisis univariado, su contribución al modelo "
  "multivariado resulta nula, por las razones de construcción del dato que se discuten más adelante.")

h2("3.5. Ingeniería de características")
p("Sobre las variables extraídas se construyeron características derivadas: edad (año T menos año "
  "de nacimiento), tiempo de afiliación, proxy de menopausia (edad ≥ 50), antecedente familiar de "
  "mama por póliza, indicadores de biopsia previa y la familia de variables PM10. Se aplicaron las "
  "siguientes decisiones, todas verificadas sobre los datos del proyecto:")
bullet("eliminación de variables duplicadas al 100% (los pares antecedente/consolidado de HTA, DM, "
       "biopsia y cáncer resultaron idénticos) y de columnas constantes tras la cohorte ≥ 18.")
bullet("winsorización a cotas de plausibilidad clínica para corregir errores de captura (por "
       "ejemplo, triglicéridos con máximo de 228.000, LDL negativo o IMC de 829).")
bullet("eliminación de PESO conservando IMC (decisión clínica: el IMC mide adiposidad sin el "
       "factor de confusión de la estatura) y de TALLA como variable independiente, rompiendo la "
       "colinealidad IMC–PESO de 0,87.")
bullet("ausencia de transformaciones logarítmicas: el análisis de gráficos cuantil-cuantil mostró "
       "que la no-normalidad provenía de la masa puntual de imputación (90–97% de los labs en la "
       "mediana) y no de colas largas, por lo que ninguna variable cumplía el criterio de cola "
       "monótona; además, para árboles las transformaciones monótonas son irrelevantes.")
bullet("codificación: ZONA y regional con one-hot; el código de municipio se descartó por aportar "
       "ruido con tan pocos positivos.")
p("Se eliminaron por fuga clínica las remisiones a mama (OR 579, presente en 78,5% de los "
  "positivos) y a oncología (OR 161): una remisión a especialidad de mama u oncología en el año T "
  "refleja una sospecha de cáncer ya existente, lo que confunde predicción prospectiva con "
  "confirmación diagnóstica. Las remisiones a ginecología y endocrinología sí se conservaron.")

h2("3.6. Variantes de imputación: mediana frente a NaN nativo")
p("El mecanismo de ausencia de los laboratorios es Missing Not At Random (MNAR): un laboratorio no "
  "ordenado no es un dato perdido al azar, sino una señal clínica en sí misma, y su ausencia "
  "alcanza el 90–97%. Por ello se descartó la imputación múltiple (MICE), que fabricaría valores y "
  "borraría la señal contenida en los indicadores {lab}_disponible. Se prepararon dos variantes "
  "con columnas idénticas que difieren solo en el tratamiento de los doce continuos imputables: "
  "una variante con mediana del train más indicadores (para modelos lineales) y otra con NaN "
  "nativo más indicadores (para los árboles, que manejan el NaN de forma nativa). Tras eliminar "
  "las características que no generalizan (el año y la frecuencia de municipio), el espacio final "
  "es de 59 características.")

h2("3.7. Control de fugas de información")
p("Se implementaron tres barreras anti-leakage. La barrera temporal restringe todas las features "
  "al año T o anteriores. La barrera de prevalencia excluye del denominador a las pacientes con "
  "diagnóstico de mama registrado en T en cualquiera de los dos sistemas, evitando confundir "
  "recurrencia con incidencia. La barrera de validación cruzada emplea GroupKFold particionado por "
  "paciente, de modo que todos los registros de una misma mujer caen en el mismo pliegue; esto es "
  "indispensable porque una paciente puede ser negativa en un par y positiva en el siguiente.")

h2("3.8. Protocolo de entrenamiento y evaluación")
p("Todos los modelos se evaluaron con GroupKFold por paciente y predicciones out-of-fold, y con "
  "una validación temporal final 2023→2024 que nunca se usó para seleccionar hiperparámetros. La "
  "regresión logística se entrenó sobre la variante imputada con escalado; los árboles sobre la "
  "variante con NaN nativo. Para el tratamiento del desbalanceo se compararon class_weight, SMOTE "
  "combinado con submuestreo (aplicados solo dentro del train de cada pliegue mediante un pipeline "
  "de imbalanced-learn) y la reponderación por scale_pos_weight. El tuning de hiperparámetros se "
  "realizó con Optuna y muestreador TPE, optimizando AUC-PR. La métrica principal es el recall, "
  "complementada por F2, AUC-PR (robusta al desbalanceo), AUC-ROC y el recall@top-k, este último "
  "directamente interpretable como la fracción de casos captados al tamizar el k% de mayor riesgo.")
p("El entorno de ejecución combina Apache PySpark para el pipeline ETL distribuido (driver 12 GB, "
  "ejecutor 6 GB, off-heap 16 GB) y un stack de Python con pandas, scikit-learn, LightGBM, XGBoost, "
  "imbalanced-learn, Optuna y SHAP. Se fijó numpy 1.26.4 para preservar la compatibilidad binaria "
  "del stack tras la instalación de Optuna y SHAP.")

h2("3.9. Infraestructura, pipeline y reproducibilidad")
p("La construcción de datos se organiza como un pipeline de cinco etapas secuenciales, cada una "
  "implementada como un artefacto independiente que lee de Google Cloud Storage y escribe el "
  "resultado procesado de vuelta al bucket, lo que permite reanudar y auditar cada paso de forma "
  "aislada:")
add_table(["Etapa", "Descripción", "Salida"],
    [["1. Radicados CAC", "Consolida post-auditoría 2021–2024 y armoniza el instructivo",
      "acumulado despues de auditoria.parquet"],
     ["2. Población", "Población mensual por año (registro de diciembre)", "{año}_seleccionadas.parquet"],
     ["3. Avicena", "Extrae variables clínicas y laboratorios", "query_avicena_{año}.parquet"],
     ["4. Unión", "Join radicados + población + Avicena", "base acumulada femenina 2021 a 2024.parquet"],
     ["5. Pares T→T+1", "Construye positivos/negativos y anonimiza", "base entrenamiento {train|validacion}.parquet"]],
    widths=[2.8, 6, 5.5], font=9)
p("La reproducibilidad se sostiene en cuatro instrumentos: scripts y notebooks versionados en Git; "
  "archivos parquet en GCS como artefactos inmutables con nombramiento semántico por año y split; "
  "archivos de parámetros en formato JSON que persisten todas las decisiones derivadas del "
  "entrenamiento (medianas, cotas de winsorización, frecuencias y conjuntos de columnas) para "
  "replicarlas en producción; y un documento de decisiones de diseño versionado junto al código, "
  "que registra las correcciones históricas y las decisiones confirmadas. Cada notebook se ejecuta "
  "de forma no interactiva sobre un kernel dedicado del entorno virtual, lo que garantiza que las "
  "versiones de librerías son las fijadas y evita dependencias accidentales del intérprete del "
  "sistema.")

h2("3.10. Justificación del diseño de validación temporal")
p("La separación entre los pares de entrenamiento (2021→2022 y 2022→2023) y el conjunto de "
  "validación (2023→2024) es prospectiva en el sentido más estricto: el modelo nunca observa el "
  "período de validación durante su ajuste. Esta es la evaluación clínicamente más relevante, pues "
  "reproduce el escenario real de despliegue, en el que un modelo entrenado con datos históricos se "
  "aplica al año inmediatamente siguiente. Una alternativa habitual —pero inadecuada aquí— sería "
  "la validación cruzada aleatoria sobre el conjunto completo, que mezclaría años y permitiría que "
  "el modelo aprendiese regularidades de un año para predecir el mismo año, inflando "
  "artificialmente el desempeño.")
p("El uso de GroupKFold por paciente en la validación cruzada interna cumple una función "
  "complementaria: garantiza que el ajuste de hiperparámetros no se beneficia de ver múltiples "
  "registros anuales de la misma mujer en distintos pliegues. Dado que una paciente puede ser "
  "negativa en un par y positiva en el siguiente, la fuga por identidad de paciente es un riesgo "
  "real que solo se neutraliza forzando a que todos sus registros caigan en el mismo pliegue. La "
  "combinación de validación temporal externa y GroupKFold interno constituye, por tanto, una "
  "doble salvaguarda contra las dos formas de fuga más probables en este problema.")

# ============================ 4. RESULTADOS Y DISCUSION ============================
h1("4. Resultados y discusión")
h2("4.1. Tasa de evento y consecuencias del desbalanceo")
p("La tasa de evento observada (0,032–0,034% según el par) es coherente con las tasas de "
  "incidencia poblacional esperadas y sitúa el problema tres órdenes de magnitud por debajo del "
  "régimen para el que se diseñaron los algoritmos estándar. La consecuencia inmediata es que la "
  "exactitud es inútil como criterio: un clasificador que prediga siempre negativo alcanza una "
  "exactitud del 99,97% con recall nulo. Por ello la evaluación se ancla en AUC-PR y recall@top-k, "
  "y el umbral de decisión se fija a posteriori en función de la capacidad del programa.")

h2("4.2. Comparativa de modelos")
p("La Tabla siguiente resume el desempeño en la validación temporal 2023→2024 de las "
  "configuraciones evaluadas a lo largo de las fases del proyecto. Los valores de AUC-PR deben "
  "leerse contra la tasa base de 0,041% de la cohorte ≥ 18: incluso los valores aparentemente "
  "bajos representan lifts de uno a dos órdenes de magnitud.")
add_table(["Modelo / configuración", "Fase", "AUC-PR", "AUC-ROC", "recall@top10%"],
    [["LightGBM-spw (spw=2307, roto)", "4", "0,0005", "0,576", "0,140"],
     ["LR-balanced", "4", "0,0020", "0,829", "0,523"],
     ["LR-balanced (tuneado)", "5b", "0,0094", "0,872", "0,634"],
     ["LR-SMOTE", "4", "0,0098", "0,879", "0,652"],
     ["LR-SMOTE (tuneado)", "5b", "0,0099", "0,878", "0,651"],
     ["XGBoost-spw (sin tunear)", "4", "0,0309", "0,881", "0,645"],
     ["LightGBM-SMOTE (tuneado)", "5", "0,0313", "0,896", "0,668"],
     ["LightGBM-SMOTE", "4", "0,0335", "0,895", "0,678"],
     ["LightGBM spw=1 (NaN nativo)", "5", "0,0346", "0,888", "0,673"],
     ["XGBoost tuneado (NaN nativo) — GANADOR", "5b", "0,0380", "0,900", "0,704"]],
    widths=[6.5, 1.5, 2, 2, 3], font=9)
figure("02_comparacion_modelos.png", "AUC-PR en validación temporal por modelo y configuración. "
       "En gris los modelos lineales, en azul las configuraciones de boosting intermedias, en "
       "verde el ex-ganador (LightGBM spw=1) y en rojo el ganador final (XGBoost tuneado).")
p("El ordenamiento es nítido: el gradient boosting domina a los modelos lineales por un factor de "
  "entre tres y diecinueve veces en AUC-PR, y dentro del boosting el XGBoost tuneado sobre datos "
  "con NaN nativo se impone con 0,0380 de AUC-PR y 0,704 de recall@top10%, superando al anterior "
  "ganador (LightGBM spw=1) en un 9,8% relativo de AUC-PR y en 3,1 puntos de recall. La Figura "
  "siguiente muestra que esta ventaja se mantiene a lo largo de toda la curva recall@top-k, lo que "
  "es relevante porque el programa operará seleccionando un percentil superior de riesgo.")
figure("03_recall_topk.png", "Recall@top-k% en validación. El XGBoost tuneado domina en todos los "
       "puntos de corte; al tamizar el 10% de mayor riesgo capta el 70,4% de los casos incidentes.")
p("La coherencia entre la validación cruzada interna (out-of-fold) y la validación temporal es la "
  "evidencia clave de que los modelos generalizan y no sobreajustan. La Tabla siguiente muestra el "
  "desempeño out-of-fold de la Fase 4, calculado sobre los pliegues GroupKFold del conjunto de "
  "entrenamiento; los valores son del mismo orden que los de la validación temporal, lo que "
  "descarta un colapso de desempeño al pasar a datos de un año posterior.")
add_table(["Modelo (CV out-of-fold)", "AUC-PR", "AUC-ROC", "recall@top5%", "recall@top10%"],
    [["LightGBM-SMOTE", "0,0268", "0,905", "0,564", "0,690"],
     ["XGBoost-spw", "0,0256", "0,891", "0,539", "0,668"],
     ["LR-SMOTE", "0,0099", "0,891", "0,507", "0,665"],
     ["LR-balanced", "0,0022", "0,823", "0,344", "0,527"],
     ["LightGBM-spw (roto)", "0,0007", "0,684", "—", "—"]],
    widths=[5.5, 2.5, 2.5, 3, 3], font=10)
p("El tuning con Optuna del SMOTE, por su parte, mostró un sobreajuste apreciable (una brecha de "
  "0,024 de AUC-PR entre el train de pliegue y la validación de pliegue) y quedó por debajo del "
  "baseline sin tunear, lo que refuerza la conclusión de que, en este problema, el resampleo "
  "sofisticado no compensa y el modelo plano sobre datos con NaN nativo es preferible.")

h2("4.3. Por qué XGBoost ganó y el papel de la reponderación")
p("El resultado metodológico más importante del proyecto es que la reponderación agresiva de "
  "clases perjudica el ordenamiento de riesgo. En la Fase 4, un LightGBM con scale_pos_weight = "
  "2.307 (igual a la razón de desbalanceo) quedó último, con AUC-PR de 0,0005 y AUC-ROC de 0,576, "
  "apenas por encima del azar. Inicialmente esto se atribuyó al NaN nativo, pero el barrido "
  "sistemático de la Fase 5 demostró que la causa era la reponderación: el AUC-PR en validación "
  "cruzada decrece de forma monótona al aumentar el scale_pos_weight, desde 0,0305 con spw = 1 "
  "hasta 0,0023 con spw = 250.")
figure("04_spw_sweep.png", "Barrido de scale_pos_weight en LightGBM (eje x logarítmico). El AUC-PR "
       "decrece monótonamente; la ausencia de reponderación (spw=1) es óptima.", 12)
p("La explicación es coherente con la teoría: la reponderación y SMOTE modifican la función de "
  "pérdida para mover la frontera de decisión, pero el AUC-PR y el recall@top-k dependen del "
  "ordenamiento de las puntuaciones, no del umbral. Al inflar artificialmente el peso de la clase "
  "positiva, el modelo sacrifica calibración del ordenamiento a cambio de mover un umbral que, en "
  "este programa, se fija de todos modos por separado mediante el percentil de riesgo. Por tanto, "
  "para una estrategia de tamización por top-k, reponderar es innecesario y contraproducente. "
  "Coherentemente, cuando se tuneó XGBoost con Optuna dejando libre el scale_pos_weight, el "
  "óptimo seleccionado fue 3,5, es decir, una reponderación leve y muy alejada del valor extremo. "
  "La comparación de Fase 4 que favorecía a SMOTE estaba confundida, pues enfrentaba un modelo con "
  "SMOTE bien configurado contra un modelo con reponderación rota.")
p("XGBoost supera a LightGBM en este conjunto por una combinación de factores: su manejo del NaN "
  "nativo se ajusta bien a la naturaleza MNAR de los laboratorios, la regularización encontrada "
  "por Optuna (gamma 3,25, min_child_weight 10,4, submuestreo 0,89 y colsample 0,87) controla el "
  "sobreajuste en un espacio de alta dimensión con señal escasa, y la reponderación leve evita el "
  "deterioro del ordenamiento observado en LightGBM. El diagnóstico out-of-fold frente a la "
  "validación confirmó la ausencia de sobreajuste.")
p("Los hiperparámetros del modelo ganador, seleccionados por Optuna optimizando AUC-PR en "
  "validación cruzada GroupKFold de tres pliegues, se documentan a continuación para su "
  "reproducibilidad:")
add_table(["Hiperparámetro", "Valor", "Hiperparámetro", "Valor"],
    [["n_estimators", "372", "gamma", "3,25"],
     ["learning_rate", "0,039", "reg_alpha", "0,79"],
     ["max_depth", "6", "reg_lambda", "1,12"],
     ["min_child_weight", "10,4", "scale_pos_weight", "3,5"],
     ["subsample", "0,89", "colsample_bytree", "0,87"]],
    widths=[4, 3, 4, 3], font=10)
p("Para acelerar la búsqueda sin comprometer la honestidad de la evaluación, el tuning se realizó "
  "sobre un submuestreo que preserva aproximadamente la prevalencia (todos los positivos más "
  "800.000 negativos), reservando el reajuste final y la validación temporal al conjunto completo. "
  "Esta práctica es legítima porque el submuestreo solo se usa para seleccionar hiperparámetros "
  "—una comparación relativa— mientras que la métrica reportada proviene siempre del conjunto de "
  "validación íntegro y no tocado.")

h2("4.4. Por qué los modelos lineales fueron insuficientes")
p("La regresión logística, aun tuneada, no superó un AUC-PR de 0,0099. La razón es estructural: el "
  "modelo lineal asume aditividad en el espacio log-odds y no captura las interacciones que "
  "gobiernan el riesgo, como la combinación de edad post-menopáusica con adiposidad elevada y "
  "antecedentes. Además, la presencia de multicolinealidad (IMC–PESO 0,87; colesterol total–LDL "
  "0,80) penaliza más a los modelos lineales que a los árboles. La regresión logística conserva su "
  "valor como referencia interpretable y como prueba de que la ganancia del boosting proviene de "
  "la no-linealidad y no de un artefacto del protocolo.")

h2("4.5. Interpretabilidad: análisis SHAP por modelo")
p("Se calcularon valores SHAP con TreeExplainer sobre una muestra compuesta por todos los "
  "positivos del train y 40.000 negativos. Las dos figuras siguientes presentan la contribución "
  "global (media del valor absoluto de SHAP) del modelo ganador (XGBoost) y del anterior ganador "
  "(LightGBM), que permiten contrastar la estabilidad de la historia de características entre "
  "algoritmos.")
figure("05_shap_xgboost.png", "Contribución global SHAP (mean|SHAP|) del XGBoost ganador. Dominan "
       "la edad y el número de consultas, seguidos a un orden de magnitud por la regional de "
       "Bogotá, el tiempo de afiliación y las variables antropométricas.", 11)
figure("06_shap_lightgbm.png", "Contribución global SHAP del LightGBM spw=1. La jerarquía es "
       "similar pero con la edad por encima del número de consultas.", 11)
p("Ambos modelos coinciden en que la edad y el número de consultas anuales son, con diferencia, "
  "las dos características más influyentes, separadas del resto por casi un orden de magnitud. "
  "Existe, sin embargo, una diferencia reveladora: en LightGBM la edad encabeza la jerarquía "
  "(mean|SHAP| 1,39 frente a 1,03 del número de consultas), mientras que en XGBoost el orden se "
  "invierte (1,37 para el número de consultas frente a 1,16 para la edad) y el indicador "
  "tiene_avicena asciende posiciones. Es decir, el modelo ganador se apoya algo más en señales de "
  "uso del sistema de salud. Los diagramas de enjambre (beeswarm) confirman además la dirección de "
  "los efectos.")
figure("11_beeswarm_xgboost.png", "Diagrama de enjambre SHAP del XGBoost ganador: dirección y "
       "magnitud del efecto de cada característica sobre la predicción individual.", 12)
figure("12_beeswarm_lightgbm.png", "Diagrama de enjambre SHAP del LightGBM spw=1, para contraste.", 12)
p("Para cuantificar con detalle la jerarquía de contribuciones, las dos tablas siguientes recogen "
  "las quince características de mayor mean|SHAP| en cada modelo, leídas directamente de los "
  "artefactos de interpretabilidad del proyecto.")
try:
    _sx = json.load(open(os.path.join(BASE, "bases", "shap_xgboost.json"), encoding="utf-8"))["mean_abs_shap"]
    _rows = [[k, f"{v:.4f}"] for k, v in list(_sx.items())[:15]]
    add_table(["Característica (XGBoost)", "mean|SHAP|"], _rows, widths=[8, 4], font=10)
    _sl = json.load(open(os.path.join(BASE, "bases", "tuning_fase5.json"), encoding="utf-8"))["shap_contribucion"]
    _rows = [[k, f"{v:.4f}"] for k, v in list(_sl.items())[:15]]
    add_table(["Característica (LightGBM)", "mean|SHAP|"], _rows, widths=[8, 4], font=10)
except Exception as _e:
    p(f"[Tablas SHAP no disponibles: {_e}]")
p("El contraste confirma la robustez de la señal principal —edad y utilización del sistema "
  "dominan en ambos— y revela que la regional de Bogotá y las variables antropométricas (IMC, "
  "perímetro abdominal, talla) ocupan posiciones intermedias estables. La menarquia y el proxy de "
  "menopausia, ambos relacionados con la exposición hormonal acumulada, aparecen igualmente en los "
  "dos rankings, lo que es clínicamente plausible y aumenta la confianza en que el modelo no se "
  "apoya únicamente en artefactos de utilización.")

h2("4.6. Por qué los indicadores de disponibilidad y los antecedentes familiares no afectan al modelo")
p("Un hallazgo consistente es que un grupo amplio de características recibe una contribución SHAP "
  "prácticamente nula. La Figura siguiente muestra las catorce de menor contribución en el modelo "
  "ganador: incluye todos los indicadores {lab}_disponible de laboratorios, las variables de "
  "cocina en recinto cerrado y fumador pasivo, y los tres antecedentes familiares "
  "(familiar_mama, familiar_cancer y familiar_cancer_mama), estos últimos con contribución "
  "exactamente cero.")
figure("07_flags_cero.png", "Características con contribución SHAP cercana a cero en el XGBoost "
       "ganador. Los antecedentes familiares aparecen con valor exactamente nulo.", 11)
p("Las razones difieren por familia de variables y conviene explicarlas con precisión. En el caso "
  "de los indicadores {lab}_disponible, su información ya está contenida, de forma redundante, en "
  "el propio tratamiento del NaN nativo del árbol y en variables de utilización más potentes como "
  "el número de consultas: el árbol aprende la dirección de ramificación para el NaN de cada "
  "laboratorio, de modo que un indicador binario adicional que codifica exactamente esa misma "
  "ausencia no añade capacidad discriminante. Dicho de otro modo, la señal de 'no tiene este "
  "laboratorio' ya está disponible por dos vías más informativas, y el indicador queda subsumido.")
p("En el caso de los antecedentes familiares, la contribución nula tiene un origen distinto y más "
  "fundamental: una limitación de construcción del dato. El antecedente familiar se aproxima "
  "vinculando a familiares dentro de la misma póliza de afiliación, lo que solo captura parentescos "
  "que también son afiliados de Sanitas y que además tienen un diagnóstico registrado. La "
  "consecuencia es que la prevalencia observada de la variable es minúscula y prácticamente no "
  "varía entre positivos y negativos; sin variabilidad informativa, el modelo no puede extraer "
  "señal alguna, por muy establecido que esté el factor de riesgo en la literatura clínica. No se "
  "trata de que la historia familiar no importe biológicamente, sino de que el proxy disponible no "
  "la mide. Variables como la cocina en recinto cerrado o el fumador pasivo combinan baja "
  "prevalencia con un efecto reducido en esta población, lo que explica su contribución marginal. "
  "Estas características son, por tanto, candidatas naturales a poda sin pérdida de desempeño.")

h2("4.7. Sesgo de vigilancia")
p("El predominio del número de consultas y de tiene_avicena entre las características más "
  "influyentes señala un sesgo de vigilancia: el modelo aprende, en parte, dónde se detecta la "
  "enfermedad —es decir, quién utiliza más el sistema de salud y por tanto tiene más "
  "oportunidades de diagnóstico— y no solo el riesgo intrínseco. Esta lógica de causalidad inversa "
  "es la misma que motivó la eliminación temprana de las remisiones a mama y oncología. El sesgo "
  "es atenuado pero no eliminado, y obliga a un análisis de sensibilidad antes de operacionalizar "
  "el modelo, así como a una vigilancia explícita de la equidad entre regiones con distinta "
  "densidad de datos clínicos.")

h2("4.8. Depuración de características y equidad geográfica")
p("A partir del análisis SHAP se evaluó la eliminación de dos características. La eliminación de "
  "tiene_avicena —definida como el indicador de tener algún registro en Avicena, y por tanto "
  "redundante con el número de consultas— no tuvo costo: el AUC-PR pasó de 0,0380 a 0,0375 y el "
  "recall@top10% de 0,704 a 0,703, por lo que se retiró del modelo y del pipeline. En cambio, la "
  "eliminación de la regional sí tuvo un costo apreciable (AUC-PR 0,0375 → 0,0313, una caída del "
  "16,5%, y recall@top10% 0,703 → 0,661) y, de manera reveladora, no resolvía la preocupación de "
  "equidad: incluso sin la variable regional, Bogotá seguía sobre-representada en el decil superior "
  "(1,20×), porque otras características actúan como proxy geográfico (mayor acceso implica más "
  "datos clínicos y más señal). Se decidió por tanto conservar la regional en el modelo y atender "
  "la equidad en la asignación de leads.")
p("La asignación estratificada —seleccionar el top-k% dentro de cada regional en lugar del top-k% "
  "global— resuelve la equidad por diseño con un costo mínimo de un punto de recall (de 0,703 a "
  "0,693). Bajo asignación estratificada, todas las regionales quedan en una sobre-representación "
  "de 1,00 (cupos proporcionales a su población). Un giro relevante es que, en la asignación "
  "global, la regional sobre-priorizada no era Bogotá (0,86×, sub-representada) sino Medellín "
  "(1,99×) y Barranquilla (1,31×). Conviene matizar que la estratificación iguala los cupos, no la "
  "cobertura: el recall por regional sigue siendo dispar (Bogotá 0,92 frente a Medellín 0,35 y "
  "Centro Oriente 0,47), lo que refleja diferencias de detectabilidad por densidad de datos y "
  "constituye una línea de mejora futura.")
p("La Tabla siguiente cuantifica el efecto de la estratificación sobre la composición de los leads. "
  "La columna de sobre-representación es el cociente entre el porcentaje de leads de la regional y "
  "su porcentaje de población; un valor de 1,00 indica reparto exactamente proporcional.")
add_table(["Regional", "Sobre-repr. (global)", "Sobre-repr. (estratificada)"],
    [["Bogotá", "0,86×", "1,00×"],
     ["Cali", "1,14×", "1,00×"],
     ["Barranquilla", "1,31×", "1,00×"],
     ["Bucaramanga", "0,92×", "1,00×"],
     ["Medellín", "1,99×", "1,00×"],
     ["Centro Oriente", "0,62×", "1,00×"]],
    widths=[4.5, 4.5, 5], font=10)
p("Mientras la asignación global produce dispersiones marcadas —Medellín casi duplica su cuota "
  "proporcional y Centro Oriente queda en 0,62×—, la estratificada las anula por construcción. El "
  "recall por regional, en cambio, permanece dispar (Bogotá 0,92, Bucaramanga 0,76, Cali 0,64, "
  "Barranquilla 0,52, Centro Oriente 0,47, Medellín 0,35), lo que evidencia que la equidad de "
  "cupos y la equidad de cobertura son objetivos distintos: la primera se resuelve en la "
  "asignación, la segunda requiere mejorar la detectabilidad del modelo en las regionales con "
  "menor densidad de datos clínicos.")

h2("4.9. Calibración y uso de las puntuaciones")
p("Para una estrategia de tamización basada en el ordenamiento por percentil de riesgo, la "
  "calibración absoluta de las probabilidades no es un requisito: basta con que el modelo ordene "
  "correctamente. No obstante, si en el futuro las puntuaciones quisieran comunicarse como "
  "'probabilidad individual de desarrollar cáncer en el próximo año' a clínicos o pacientes, sería "
  "imprescindible una calibración posterior (isotónica o de Platt), porque bajo desbalanceo "
  "extremo y con reponderación las probabilidades de salida están sistemáticamente sesgadas. La "
  "recomendación operativa es, por tanto, usar el modelo como generador de un ranking de riesgo y "
  "no como estimador de probabilidad absoluta hasta que se complete un estudio de calibración.")

h2("4.10. Amenazas a la validez")
p("Se reconocen varias amenazas a la validez. En cuanto a la validez interna, el sesgo de "
  "vigilancia descrito implica que parte de la señal proviene de la utilización del sistema y no "
  "del riesgo intrínseco; aunque se mitigó eliminando las remisiones a mama y oncología y la "
  "variable tiene_avicena, el número de consultas permanece como predictor dominante. Respecto a "
  "la validez de construcción, la etiqueta depende de la radicación ante la CAC, de modo que un "
  "caso incidente real no radicado o radicado con retraso sería tratado como negativo, "
  "introduciendo ruido en la clase minoritaria. En cuanto a la validez externa, el cambio en el "
  "tamaño de la población entre pares —los excluidos por falta de continuidad pasan de 16.026 en "
  "2021→2022 a 271.226 en 2023→2024— refleja fluctuaciones reales de afiliación, posiblemente "
  "asociadas a efectos post-pandemia, que hacen que los conjuntos no sean estrictamente comparables "
  "en distribución. Finalmente, la limitación del antecedente familiar (solo afiliados Sanitas con "
  "diagnóstico registrado) y la extrapolación del PM10 para 2021–2022 son fuentes conocidas de "
  "error que el modelo gestiona, en el segundo caso, mediante el indicador de disponibilidad.")

h2("4.11. Consideraciones éticas")
p("El despliegue de un modelo predictivo de riesgo en salud exige salvaguardas explícitas. La "
  "anonimización de los identificadores mediante hash SHA-256 protege la identidad de las "
  "afiliadas en los artefactos de modelado. La decisión de conservar la regional y resolver la "
  "equidad en la asignación estratificada responde a un principio de justicia distributiva: evitar "
  "que la mayor densidad de datos de unas regiones se traduzca en una concentración desproporcionada "
  "de la oferta de tamización. El uso de SHAP responde al principio de transparencia y al requisito "
  "de explicabilidad para la validación por comités clínicos. Es importante subrayar que el modelo "
  "no sustituye el juicio clínico ni los criterios de tamización por edad establecidos, sino que "
  "los complementa orientando recursos adicionales hacia el grupo donde el cribado rutinario no "
  "llega; ninguna paciente queda excluida de la atención estándar por efecto del modelo.")

# ============================ 5. DISENO OPERATIVO ============================
h1("5. Diseño operativo y ensayo A/B")
h2("5.1. Punto de operación")
p("Sobre el modelo ganador se construyó la curva de operación del programa, que para cada "
  "percentil de riesgo reporta los cupos (mujeres a tamizar), el recall, el valor predictivo "
  "positivo, el lift sobre la tasa base y el número necesario a tamizar (NNS) para hallar un caso. "
  "La validación contiene 2.410.807 mujeres y 993 casos, con tasa base de 0,041%.")
add_table(["Top-k%", "Cupos", "Recall", "VPP", "Lift", "NNS"],
    [["0,5%", "12.052", "26,6%", "2,19%", "53×", "46"],
     ["1%", "24.105", "32,3%", "1,33%", "32×", "75"],
     ["2%", "48.214", "42,0%", "0,86%", "21×", "116"],
     ["5%", "120.538", "55,0%", "0,45%", "11×", "221"],
     ["10%", "241.078", "69,3%", "0,29%", "6,9×", "350"],
     ["20%", "482.160", "80,8%", "0,17%", "4,0×", "601"]],
    widths=[2.5, 3, 2.5, 2.5, 2, 2.5], font=10)
figure("08_curva_operacion.png", "Curva de operación (asignación estratificada): recall (verde) y "
       "número necesario a tamizar (rojo) en función del percentil de riesgo seleccionado.", 12)
p("La lectura del trade-off es directa: cortes pequeños maximizan la eficiencia (lift alto, NNS "
  "bajo) a costa de cobertura, mientras que cortes amplios elevan la cobertura disparando el "
  "volumen de tamizaciones. El codo de la curva se sitúa entre el 1% y el 5%.")

h2("5.2. Estructura por edad y rol incremental del modelo")
p("La restricción operativa relevante es que el programa garantiza tamizaje bienal al menos en "
  "mujeres mayores de 45 años, que es el estándar de atención. En consecuencia, el modelo no "
  "necesita racionar a ese grupo —ya cubierto— y su valor incremental se concentra en la "
  "subpoblación menor de 45 años, donde el cribado por edad no llega. El análisis por bandas de "
  "edad cuantifica este punto.")
add_table(["Banda", "Mujeres", "% casos", "Tasa banda"],
    [["18–39", "1.127.420", "10,8%", "0,009%"],
     ["40–44", "253.614", "12,4%", "0,048%"],
     ["45–49", "204.586", "11,3%", "0,055%"],
     ["50–59", "348.325", "23,0%", "0,065%"],
     ["60+", "476.862", "42,6%", "0,089%"]],
    widths=[3, 3.5, 3, 3.5], font=10)
figure("09_estructura_edad.png", "Estructura por edad de los casos incidentes: porcentaje de casos "
       "(barras) y tasa de incidencia por banda (línea). La línea discontinua marca el corte de 45 "
       "años del tamizaje rutinario.", 12)
p("Los casos menores de 45 años suman 230 (23,2% del total), de los cuales la mayoría se concentra "
  "en la banda 40–44, con una tasa cinco veces superior a la de 18–39. El tamizaje bienal ≥ 45 ya "
  "cubre por edad cerca del 77% de los casos, de modo que el aporte inequívoco del modelo es "
  "identificar el alto riesgo entre las menores de 45 para una tamización proactiva temprana. "
  "Dentro de esta subpoblación, seleccionar el 2% de mayor riesgo (27.620 mujeres) capta el 54,8% "
  "de los casos menores de 45, con un NNS de 219.")

h2("5.3. Diseño y dimensionamiento del ensayo A/B")
p("Se definió un ensayo A/B de brazo único sobre la subpoblación menor de 45 años de alto riesgo: "
  "el grupo de intervención recibe mamografía proactiva y el de control la atención estándar (el "
  "cribado se inicia a los 45). El endpoint primario es el estadio al diagnóstico (temprano I–IIa "
  "frente a tardío III–IV); la línea base de estadio temprano se toma del 54% reportado por CAC "
  "(2024). Como la incidencia es rara, el número de casos acumulados por brazo condiciona la "
  "potencia. La Figura siguiente muestra el tamaño de muestra requerido por brazo según el efecto "
  "a detectar, con potencia del 80% y α = 0,05 bilateral.")
figure("10_poder_ab.png", "Casos confirmados requeridos por brazo (escala logarítmica) en función "
       "del porcentaje de estadio temprano alcanzado por la intervención, partiendo de un 54% en "
       "el grupo de control.", 12)
add_table(["Estadio temprano intervención", "Δ (pp)", "Casos/brazo", "Casos total"],
    [["62%", "+8", "597", "1.194"],
     ["66%", "+12", "261", "522"],
     ["70%", "+16", "144", "288"],
     ["74%", "+20", "90", "180"]],
    widths=[5.5, 2.5, 3, 3], font=10)
p("Con el corte recomendado del 2% de la subpoblación menor de 45 (≈ 27.620 mujeres por año, "
  "asignadas 1:1 a nivel individual, lo que acumula unos 63 casos por brazo y año), detectar un "
  "incremento de 16 a 20 puntos porcentuales en estadio temprano requiere entre 1,4 y 2,3 años de "
  "reclutamiento. La randomización individual maximiza la potencia y evita el efecto de diseño de "
  "una aleatorización por conglomerados; esta última solo se justificaría ante riesgo de "
  "contaminación operativa entre brazos. Como endpoint secundario, más potente y de lectura más "
  "temprana, se propone la tasa de detección confirmada por brazo, que usa todos los leads como "
  "denominador.")
p("Procede señalar tres salvedades sobre el dimensionamiento. La línea base del 54% es poblacional "
  "y no específica de las menores de 45; dado que los cánceres en mujeres jóvenes tienden a ser "
  "más agresivos (mayor proporción de triple negativo), la base de estadio temprano real en esta "
  "subpoblación podría ser menor, por lo que conviene estimarla de las bases internas antes de "
  "fijar el efecto. El tamaño del efecto a detectar es un supuesto del analista y no un dato. Y el "
  "valor predictivo positivo empleado proviene de un horizonte de un año, por lo que la "
  "acumulación multianual asume su estabilidad.")

h2("5.4. Evolución del desempeño a lo largo del proyecto")
p("El desempeño del mejor modelo evolucionó de forma instructiva a través de las fases, y su "
  "trazabilidad ilustra cómo decisiones aparentemente menores tienen efectos mayores. En la Fase 4 "
  "el ganador fue LightGBM con SMOTE (AUC-PR 0,0335), mientras que el LightGBM con reponderación "
  "extrema quedó último por el error de configuración del scale_pos_weight. La Fase 5 corrigió ese "
  "diagnóstico: al barrer el scale_pos_weight se descubrió que el problema no era el NaN nativo sino "
  "la reponderación, y un LightGBM plano con spw=1 sobre datos nativos pasó a liderar (AUC-PR "
  "0,0346). La Fase 5b tuneó los modelos previamente perdedores y reveló que un XGBoost con "
  "reponderación leve superaba a todos (AUC-PR 0,0380). Finalmente, la depuración de características "
  "guiada por SHAP retiró tiene_avicena sin costo (AUC-PR 0,0375) y descartó la eliminación de la "
  "regional por su costo y su nulo beneficio de equidad. El recorrido total representa una mejora "
  "del 13% en AUC-PR respecto al primer ganador, obtenida no por mayor complejidad sino por "
  "corregir supuestos y depurar el espacio de características.")
add_table(["Fase", "Modelo ganador de la fase", "AUC-PR", "Aprendizaje clave"],
    [["4", "LightGBM-SMOTE", "0,0335", "SMOTE aparentaba ser necesario"],
     ["5", "LightGBM spw=1 (nativo)", "0,0346", "La reponderación extrema era el problema, no el NaN"],
     ["5b", "XGBoost tuneado (nativo)", "0,0380", "Reponderación leve + regularización ganan"],
     ["6", "XGBoost sin tiene_avicena", "0,0375", "Poda sin costo; equidad en la asignación"]],
    widths=[1.5, 6, 2, 5], font=9)

# ============================ 6. CONCLUSIONES ============================
h1("6. Conclusiones")
p("Primera. Es viable construir, sobre fuentes administrativas y clínicas de escala poblacional, "
  "un modelo de predicción prospectiva de incidencia de cáncer de mama con control estricto de "
  "fugas de información. El diseño de pares T→T+1, la exclusión de prevalentes y la validación "
  "cruzada por paciente conforman un protocolo reproducible cuya validación temporal 2023→2024 "
  "simula fielmente el despliegue real. La principal limitación es el desbalanceo extremo "
  "(≈ 3.000:1), inherente a un problema de incidencia y no a un defecto del diseño.")
p("Segunda. El gradient boosting supera de forma clara a los modelos lineales, y entre las "
  "configuraciones de boosting el XGBoost tuneado sobre datos con NaN nativo es el mejor modelo "
  "(AUC-PR 0,0380; recall@top10% 0,704; AUC-ROC 0,900). La implicación práctica es que, al tamizar "
  "el decil superior de riesgo, el programa capturaría alrededor del 70% de los casos incidentes, "
  "un lift cercano a siete veces sobre la selección aleatoria. La mejora adicional pasa por "
  "disponer de mejores características, no por más tuning, pues el desempeño se encuentra saturado "
  "en torno a 0,034–0,038 de AUC-PR.")
p("Tercera. La reponderación agresiva de clases degrada el ordenamiento de riesgo y, por tanto, el "
  "AUC-PR y el recall@top-k; el óptimo es la reponderación nula o leve, fijando el punto de corte "
  "por separado mediante el percentil de riesgo. Este resultado contradice la intuición habitual de "
  "'compensar' el desbalanceo con pesos extremos y es transferible a otros problemas de tamización "
  "por ranking. Como limitación, este hallazgo se estableció para la métrica de ordenamiento y no "
  "necesariamente aplica a escenarios donde importe la calibración absoluta de la probabilidad.")
p("Cuarta. La interpretabilidad SHAP no solo explica el modelo sino que orienta decisiones de "
  "diseño: confirmó el sesgo de vigilancia, justificó la poda de los indicadores de disponibilidad "
  "y de los antecedentes familiares —estos últimos sin señal por una limitación de construcción "
  "del proxy, no por irrelevancia biológica— y motivó conservar la regional resolviendo la equidad "
  "en la asignación estratificada. Como trabajo futuro se plantea estimar la línea base de estadio "
  "específica de la subpoblación menor de 45 años, mejorar la detectabilidad en regionales con "
  "menor densidad de datos, evaluar la incorporación del PM10 y ejecutar el ensayo A/B con el "
  "diseño dimensionado.")
p("Quinta. El modelo se traduce en un programa operativo concreto y defendible. La curva de "
  "operación permite fijar el punto de corte en función de la capacidad instalada, y el análisis "
  "por edad redefine el rol del modelo como complemento —no sustituto— del tamizaje bienal por "
  "edad: su valor incremental se concentra en la subpoblación menor de 45 años, donde el cribado "
  "rutinario no llega y se concentra el 23,2% de los casos. El ensayo A/B propuesto, de brazo "
  "único y randomización individual sobre el 2% de mayor riesgo de esa subpoblación, está "
  "dimensionado para detectar un desplazamiento de estadio clínicamente significativo en un "
  "horizonte de uno a dos años. La principal cautela es que la línea base de estadio empleada es "
  "poblacional y deberá estimarse de forma específica para las menores de 45 antes de cerrar el "
  "protocolo, dado que los tumores en mujeres jóvenes tienden a presentaciones más agresivas.")

# ============================ 7. REFERENCIAS ============================
h1("7. Referencias")
refs = [
 "AlGhunaim, S., & Al-Baity, H. H. (2019). On the scalability of machine-learning algorithms for "
 "breast cancer prediction in big data context. IEEE Access, 7, 91535–91546.",
 "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic Minority "
 "Over-sampling Technique. Journal of Artificial Intelligence Research, 16, 321–357.",
 "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the "
 "22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785–794.",
 "Cuenta de Alto Costo (CAC). (2024). Situación del cáncer en la población adulta atendida en el "
 "SGSSS de Colombia. Fondo Colombiano de Enfermedades de Alto Costo.",
 "Gail, M. H., Brinton, L. A., Byar, D. P., Corle, D. K., Green, S. B., Schairer, C., & Mulvihill, "
 "J. J. (1989). Projecting individualized probabilities of developing breast cancer for white "
 "females. Journal of the National Cancer Institute, 81(24), 1879–1886.",
 "Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T.-Y. (2017). LightGBM: "
 "A highly efficient gradient boosting decision tree. Advances in Neural Information Processing "
 "Systems, 30, 3146–3154.",
 "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. "
 "Advances in Neural Information Processing Systems, 30, 4765–4774.",
 "Mariotto, A. B., Yabroff, K. R., Shao, Y., Feuer, E. J., & Brown, M. L. (2011). Projections of "
 "the cost of cancer care in the United States: 2010–2020. Journal of the National Cancer "
 "Institute, 103(2), 117–128.",
 "Singh, A., & Dubey, S. (2024). Machine learning approaches for breast cancer prediction: A review.",
 "Sung, H., Ferlay, J., Siegel, R. L., Laversanne, M., Soerjomataram, I., Jemal, A., & Bray, F. "
 "(2021). Global cancer statistics 2020: GLOBOCAN estimates of incidence and mortality worldwide "
 "for 36 cancers in 185 countries. CA: A Cancer Journal for Clinicians, 71(3), 209–249.",
 "Tahir, M., & Khan, A. (2025). Deep learning for mammographic breast cancer detection: A "
 "systematic review.",
 "Tyrer, J., Duffy, S. W., & Cuzick, J. (2004). A breast cancer prediction model incorporating "
 "familial and personal risk factors. Statistics in Medicine, 23(7), 1111–1130.",
 "Waks, A. G., & Winer, E. P. (2019). Breast cancer treatment: A review. JAMA, 321(3), 288–300.",
 "Yavuz, E., et al. (2023). Comparison of machine-learning algorithms for breast cancer "
 "classification on balanced and imbalanced datasets.",
]
for r in refs:
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.left_indent = Cm(1); par.paragraph_format.first_line_indent = Cm(-1)
    run = par.add_run(r); run.font.name = "Times New Roman"; run.font.size = Pt(11)

# ============================ 8. CODIGO FUENTE Y DATOS ============================
h1("8. Código fuente y datos analizados")
p("El código fuente desarrollado durante este Trabajo de Fin de Estudios se encuentra alojado en "
  "el siguiente repositorio público de GitHub:")
par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = par.add_run("https://github.com/jucasamesa/comparativa_modelos_prediccion_mama")
run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(12)
blank()
p("El estudiante es el único autor del código y único propietario del repositorio. En el "
  "repositorio no existe ningún commit de otro usuario distinto del autor.")
p("Este Trabajo de Fin de Estudios está asociado a un proyecto de la Cuenta de Alto Costo de EPS "
  "Sanitas. Por razones de confidencialidad de la información clínica de las pacientes, y conforme "
  "a la normativa de protección de datos de salud, no se ponen a disposición pública ni las bases "
  "de datos utilizadas ni la parte del código que accede a la infraestructura interna de la "
  "entidad. En concreto, quedan excluidos del repositorio: (1) las fuentes de datos empleadas en "
  "el análisis —los radicados de la Cuenta de Alto Costo, la población de afiliados y la historia "
  "clínica electrónica Avicena—, que contienen información sensible de salud aun después de la "
  "anonimización; y (2) el pipeline de construcción de datos implementado en PySpark, que se "
  "conecta a los buckets de almacenamiento en la nube y a las consultas internas de Avicena y que, "
  "por tanto, es directamente asociable a la entidad.")
p("El repositorio contiene exclusivamente el código de modelado, comparación de algoritmos, "
  "interpretabilidad y diseño operativo que opera sobre bases de datos anonimizadas —en las que el "
  "identificador de paciente fue transformado mediante los primeros dieciséis caracteres de un "
  "hash SHA-256—, junto con los artefactos de resultados agregados (métricas de validación, "
  "valores de contribución SHAP, curvas de operación y dimensionamiento del ensayo) y las figuras "
  "incluidas en esta memoria. Ninguno de estos artefactos contiene información a nivel de paciente. "
  "De este modo se preserva la reproducibilidad de los análisis y figuras del presente documento "
  "sin comprometer la confidencialidad de los datos de origen.")

doc.save(OUT_PATH)
print("Documento generado:", OUT_PATH)
